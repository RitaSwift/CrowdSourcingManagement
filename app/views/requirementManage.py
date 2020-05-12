from flask import Blueprint, render_template, request, make_response, send_file
from app.models import Project, Requirement, db
from app.nlp_api.doc2Similarity import _d2v_model, _ds2v_vector, _manhattan_sim, _euclidean_sim
from app.nlp_api.doc2vec_sim_nff1026 import *
from app.nlp_api.GRU_w2v_glove_predict import GRU_Glove
from app.nlp_api.priority import confirm_priority
from docx import Document
from docx.oxml.ns import qn
import os
requirementManage = Blueprint('requirementManage', __name__, url_prefix='/')

@requirementManage.route('/rm')
def rm():
    pid = int(request.args.get('pid'))
    project = Project.query.filter_by(pid=pid).first()
    requirementList = Requirement.query.filter_by(pid=pid).all()
    return render_template('requireManagement.html', project=project, requirementList=requirementList)


# @requirementManage.route('/reorganizeRequirement', methods=['GET', 'POST'])
# def rm_repeated_requirement():
#     pid = int(request.args.get('pid'))
#     if not os.path.exists('doc2vec_model'):
#         train_d2v_model("./app/nlp_api/keepass.txt")
#     requirements = Requirement.query.filter_by(pid=pid).all()
#     similar = {}
#     for i in range(0, len(requirements)):
#         for j in range(i + 1, len(requirements)):
#             vector1 = d2v_vector(list(requirements[i].description))
#             vector2 = d2v_vector(list(requirements[j].description))
#             if (manhattan_sim(vector1, vector2) + euclidean_sim(vector1, vector2)) / 2 > 0.90:  # 计算相似度是否超过预设值
#                 if requirements[i].rid > requirements[j].rid:
#                     similar[requirements[i].rid] = requirements[j].rid    # 将相似的两个需求关联起来，rid 大的需求指向 rid 小的需求
#                 else:
#                     similar[requirements[j].rid] = requirements[i].rid
#     # 运用并查集算法，将所有相似的需求都指向 rid 最小的那个
#     for requirement in requirements:
#         rid = requirement.rid
#         while rid in similar.keys() and similar[rid] != rid:
#             rid = similar[rid]
#         similar[requirement.rid] = rid
#         # similar[rid] = rid
#     print(similar)
#     for rid in similar.keys():
#         Requirement.query.filter_by(rid=rid).update({
#             "similar": similar[rid]
#         })
#     db.session.commit()
#     return "success"

@requirementManage.route('/reorganizeRequirement', methods=['GET', 'POST'])
def rm_repeated_requirement():
    pid = int(request.args.get('pid'))
    if not os.path.exists('doc2vec_model'):
        _d2v_model("./app/nlp_api/keepass.txt")
    requirements = Requirement.query.filter_by(pid=pid).all()
    requires = []
    for requirement in requirements:
        description = requirement.description.strip().split(' ')
        requires.append(description)
    similar = {}
    vecs = _ds2v_vector(requires, "doc2vec_model")
    for i, vec in enumerate(vecs):
        result = (_manhattan_sim(vec, vecs) + _euclidean_sim(vec, vecs)) / 2
        index = np.argwhere(result > 0.88)
        for j in index.reshape(-1):
            if i == j:
                continue
            if requirements[i].rid > requirements[j].rid:
                similar[requirements[i].rid] = requirements[j].rid  # 将相似的两个需求关联起来，rid 大的需求指向 rid 小的需求
            else:
                similar[requirements[j].rid] = requirements[i].rid

                # for i in range(0, len(requirements)):
                #     for j in range(i + 1, len(requirements)):
                #         vector1 = _d2v_vector(list(requirements[i].description))
                #         vector2 = _d2v_vector(list(requirements[j].description))
                #         if (_manhattan_sim(vector1, vector2) + _euclidean_sim(vector1, vector2)) / 2 > 0.90:  # 计算相似度是否超过预设值
                # 运用并查集算法，将所有相似的需求都指向 rid 最小的那个

    for requirement in requirements:
        rid = requirement.rid
        while rid in similar.keys() and similar[rid] != rid:
            rid = similar[rid]
        similar[requirement.rid] = rid
        # similar[rid] = rid
    # print("similar" + str(similar))
    for rid in similar.keys():
        Requirement.query.filter_by(rid=rid).update({
            "similar": similar[rid]
        })
    db.session.commit()
    return "success"


@requirementManage.route('/repeat')
def repeat():
    pid = request.args.get('pid')
    requirements = Requirement.query.filter_by(pid=pid).all()
    requirementList = {}
    # print(requirements)
    for r in requirements:
        if r.similar not in requirementList:
            requirementList[r.similar] = []
        requirementList[r.similar].append(r)
    requirements = list(requirementList.values())
    return render_template('/repeatManagement.html', requirementList=requirements)



# 分类和优先级的界面
@requirementManage.route('/priorityAndClass', methods=['GET', 'POST'])
def rm_priority_and_class():
    pid = request.args.get('pid')
    page = int(request.args.get('page')) or 1
    requirements = Requirement.query.filter_by(pid=pid, selected=0).all()
    project = Project.query.filter_by(pid=pid).first()
    pageNumber = len(requirements) // 10 + 1
    requirementList = requirements[(page - 1) * 10: page * 10]
    descriptions = []
    for requirement in requirements:
        descriptions.append(requirement.description)
    priority = confirm_priority(descriptions)  # 优先级预测
    category = GRU_Glove(descriptions)  # 进行分类
    for index, value in enumerate(requirements):
        value.rtype = category[index]
        value.priority = int(priority[index])
        Requirement.query.filter_by(rid=value.rid).update({
            'rtype': value.rtype,
            'priority': value.priority
        })
    db.session.commit()
    # requires = []
    # for requirement in requirements:
    #     description = requirement.description.strip().split(' ')
    #     requires.append(description)
    # Requirement.query.filter_by(pid=pid).update({
    #     "selected": 1
    # })
    return render_template('priorityAndClass.html', project=project, requirementList=requirementList, pageNumber=pageNumber, numOfRequirement=len(requirements))


@requirementManage.route('/selectRequirement', methods=['POST'])
def update_select_requirement():
    pid = request.args.get('pid')
    data = request.get_json()
    selected_rid = data['selected']
    for rid in selected_rid:
        Requirement.query.filter_by(rid=int(rid)).update({
            'selected': 0
        })
    db.session.commit()
    return 'success'


@requirementManage.route('/downloadDoc', methods=['GET'])
def down_rm_doc():
    pid = request.args.get('pid')
    if not pid:
        return 'false'
    else:
        PATH = 'app/static/upload/'
        project = Project.query.get(int(pid))
        requirements = Requirement.query.filter_by(pid=pid, selected=0).all()
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.add_heading(project.pname, level=1)
        p = document.add_paragraph()
        p.add_run('项目发布者 ').bold = True
        p.add_run(project.ppublisher + '\t')
        p.add_run('项目语言 ').bold = True
        p.add_run(project.planguage + '\t')
        p.add_run('项目领域 ').bold = True
        p.add_run(project.pfield + '\t')
        p = document.add_paragraph()
        p.add_run(project.pdescription)
        document.add_paragraph()
        document.add_heading('需求列表', level=2)
        for requirement in requirements:
            document.add_heading(requirement.rname, level=3)
            p = document.add_paragraph()
            p.add_run('需求提供者 ').bold = True
            p.add_run(requirement.uid + '\t')
            p.add_run('需求类型 ').bold = True
            p.add_run(requirement.rtype + '\t')
            p = document.add_paragraph()
            p.add_run('优先级 ').bold = True
            p.add_run(str(requirement.priority) + '\t')
            p.add_run('提供时间 ').bold = True
            p.add_run(str(requirement.posttime) + '\t')
            document.add_paragraph(requirement.description)

        filename = project.pname + '.docx'
        document.save(PATH + filename)
        response = make_response(send_file('static/upload/' + filename))
        response.headers["Content-Disposition"] = "attachment; filename=Requirement.docx;"
    return response
